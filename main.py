# from API_KEY import API_KEY
import streamlit as st
import openai as oai
from openai import OpenAI
import time
import random

# oai.api_key = API_KEY

api_key = st.secrets.api_key

selected_options = {}
if 'dic' not in st.session_state:
    st.session_state.dic = selected_options

st.title("Birdhaus Book Writing Online Platform Powered by AI")
# mode_topic = None
st.sidebar.image('logo.jpg')
mode_topic = st.sidebar.selectbox('Topic', ['IT', 'Marketing', 'Medicine', 'Sport', 'Culture'], index=None)
st.session_state.dic['TOPIC'] = mode_topic

select_subtopic = None
if mode_topic == 'Marketing':
    select_subtopic = st.sidebar.selectbox('Marketing', ['Traditional', 'Internet', 'Interactive'])
    st.session_state.dic['SUBTOPIC'] = select_subtopic
else:
    st.sidebar.write('Marketing not selected :(')

st.session_state.dic['DOMAIN'] = None
if select_subtopic == 'Internet':
    select_domain = st.sidebar.radio('Select a domain of your topic',
                            ['Search Engine Marketing', 'Social Media Marketing', 'Email Marketing',
                             'Content Marketing', 'Affiliate Marketing', 'Video Marketing', 'Mobile Marketing'],
                                     index=None
    )
    st.session_state.dic['DOMAIN'] = select_domain

mode_brand = st.sidebar.selectbox('Brand', ['values', 'vision', 'voice', 'visuals'])


mode_audience_age = st.sidebar.selectbox('Audience: AGE',
                                 ['18-25', '26-40', '41-60', '61+'], index=None)
st.session_state.dic['AGE'] = mode_audience_age

mode_audience_exp = st.sidebar.selectbox('Audience: Expertise',
                                     ['switcher', 'newcomer', 'average', 'expert'], index=None)
st.session_state.dic['Expertise'] = mode_audience_exp

mode_audience_resp = st.sidebar.selectbox('Audience: Responsibility',
                                     ['operation level', 'tactic level', 'strategic level', 'top level'], index=None)
st.session_state.dic['Responsibility'] = mode_audience_resp

mode_style = st.sidebar.selectbox('Text style', ['scientific', 'official business', 'journalistic', 'art', 'colloquial'])
st.session_state.dic['Text style'] = mode_style

options = st.sidebar.button('Show options')

if options:
    st.write('Selected options: ', st.session_state.dic)

def chat_gpt(prompt, tokens=None, temperature=None):
    client = OpenAI(api_key=api_key)
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content.strip()

# print(selected_options)
title_book = st.sidebar.button('Generate Title book using AI')

title = None
# st.session_state.dic['Title'] = title
if title_book:
    prompt = f"Specify 3 titles for couch book with this parameters :\n {st.session_state.dic}." \
                     f"Don't specify age in title but use linguistic term"
    client = OpenAI(api_key=oai.api_key)
    gen_titles = chat_gpt(prompt)

    st.write(gen_titles)
    titles = gen_titles.split('\n')
    # print(titles)
    title = st.selectbox('Select one title', titles)
    st.session_state.dic['Title'] = title[3:]

contents_button = st.sidebar.button('Generate Table of contents using AI')
if contents_button:
    prompt_contents = f"Specify table of contents for couch book with this parameters :\n {st.session_state.dic}." \
                      f"print the title first and then the contents"
    client = OpenAI(api_key=oai.api_key)
    gen_content = chat_gpt(prompt_contents)

    contents = gen_content.split('\n')
    print(contents)
    contents.remove('')
    st.session_state.dic['contents'] = contents
    for i in contents:
        st.write(i)
    st.balloons()

print(f'opts: {st.session_state.dic}')

import docx
from io import BytesIO
docx_binary = None
dwnld_button = None

if contents_button:
    title_contents = docx.Document()
    for i in st.session_state.dic['contents']:
        title_contents.add_paragraph(i)
    fname = f'title_content.docx'
    # title_contents.save(name)
    docx_binary = BytesIO()
    title_contents.save(docx_binary)
    docx_binary.seek(0)
    dwnld_button = st.download_button('Download file', docx_binary, file_name=fname,
                       mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

chapter_button = st.sidebar.button('Generate Chapter content using AI')
if 'chapter' not in st.session_state:
    st.session_state.chapter = 'some chapter name'

def callback1():
    st.header('Thank you for choosing our service!!!', divider='rainbow')
    st.header('BBWOP is :blue[cool] :sunglasses:')

if chapter_button:
    chapters_dict = {k: v for k, v in enumerate(st.session_state.dic['contents'])}
    num_chapter = random.choice(range(2, len(st.session_state.dic['contents'])))
    # print(chapters_dict)
    # print(num_chapter)
    st.session_state.chapter = st.selectbox('Select chapter', st.session_state.dic['contents'], num_chapter)
    # print(st.session_state.chapter)

    st.write('Selected chapter', st.session_state.chapter)
    prompt_chapter = f"You as a professional in {st.session_state.dic['DOMAIN'] if st.session_state.dic['DOMAIN'] else st.session_state.dic['TOPIC']} " \
                         f"write please a 2-page draft of a section name: ' {st.session_state.chapter} ' " \
                         f"for a book with these parameters: \n {st.session_state.dic}\n" \
                         f"print only section name and content"

    client = OpenAI(api_key=oai.api_key)
    gen_chapter = chat_gpt(prompt_chapter)
    st.text_area(f'Chapter {chapters_dict[num_chapter]}', gen_chapter)
    st.balloons()
    contents_chapter = docx.Document()
    contents_chapter.add_paragraph(gen_chapter)
    fname2 = f'Chapter_{st.session_state.chapter.strip()[:3]}.docx'
    docx_binary2 = BytesIO()
    contents_chapter.save(docx_binary2)
    docx_binary2.seek(0)
    dwnld_button2 = st.download_button('Download file', docx_binary2, file_name=fname2,
                           mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                       on_click=callback1)

